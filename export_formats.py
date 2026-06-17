"""Convert extracted Q&A content to .txt and .docx."""

from __future__ import annotations

import re
from io import BytesIO

SECTION_RE = re.compile(r"^\\section\*\{([^}]+)\}$")
TITLE_RE = re.compile(r"^\\title\{\s*\n?(.+?)\n?\}$", re.DOTALL)


def _clean_latex_inline(text: str) -> str:
    text = re.sub(r"\$\$([^$]+)\$\$", r"\1", text)
    text = re.sub(r"\$([^$]+)\$", r"\1", text)
    text = re.sub(r"\\mathrm\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\text\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\rightarrow", "->", text)
    text = re.sub(r"\\longrightarrow", "->", text)
    text = re.sub(r"\\xrightarrow\{[^}]*\}", "", text)
    text = re.sub(r"\\uparrow", "^", text)
    text = re.sub(r"\\downarrow", "v", text)
    text = re.sub(r"\\_", "_", text)
    text = re.sub(r"\\\s+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def split_blocks(content: str) -> list[str]:
    return [b.strip() for b in content.split("\n\n") if b.strip()]


def to_plain_text(content: str) -> str:
    parts: list[str] = []
    for block in split_blocks(content):
        m_title = TITLE_RE.match(block)
        if m_title:
            title = _clean_latex_inline(m_title.group(1))
            parts.append(title)
            parts.append("=" * min(60, max(len(title), 20)))
            continue

        lines = block.splitlines()
        first = lines[0].strip() if lines else ""
        m_sec = SECTION_RE.match(first)
        if m_sec:
            parts.append("")
            parts.append(m_sec.group(1).upper())
            parts.append("-" * len(m_sec.group(1)))
            for line in lines[1:]:
                if line.strip():
                    parts.append(_clean_latex_inline(line))
            continue

        for line in lines:
            if line.strip():
                parts.append(_clean_latex_inline(line))

    return "\n".join(parts).strip() + "\n"


def to_docx_bytes(content: str, title: str = "") -> bytes:
    from docx import Document

    doc = Document()
    if title:
        doc.add_heading(_clean_latex_inline(title), level=0)

    for block in split_blocks(content):
        m_title = TITLE_RE.match(block)
        if m_title:
            if not title:
                doc.add_heading(_clean_latex_inline(m_title.group(1)), level=0)
            continue

        lines = block.splitlines()
        first = lines[0].strip() if lines else ""
        m_sec = SECTION_RE.match(first)
        if m_sec:
            doc.add_heading(m_sec.group(1), level=1)
            for line in lines[1:]:
                if line.strip():
                    doc.add_paragraph(_clean_latex_inline(line))
            continue

        for line in lines:
            if line.strip():
                doc.add_paragraph(_clean_latex_inline(line))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def output_filename(stem: str, ext: str) -> str:
    return f"{stem}_qa.{ext.lstrip('.')}"
