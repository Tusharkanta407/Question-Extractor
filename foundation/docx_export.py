"""Export Foundation document to .docx format."""

from __future__ import annotations

import random
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from foundation.models import FoundationDocument, FoundationQuestion, QuestionOption

_NONNEG_NUMERIC_RE = re.compile(r"^\+?\d+(\.\d+)?$")


def _is_nonneg_numeric(value: str) -> bool:
    return bool(_NONNEG_NUMERIC_RE.match(value.strip()))


def _build_fib_pool(doc: FoundationDocument) -> list[str]:
    """Collect all FIB answers (word and numeric) to use as distractor pool for SCQ options."""
    return [
        q.answer_key.strip()
        for q in doc.questions
        if q.question_type == "FIB"
        and (q.answer_key or "").strip()
    ]


def _build_fib_options(correct: str, pool: list[str]) -> tuple[list[QuestionOption], str]:
    """Build 4 MCQ options. Correct answer placed randomly in (a-d)."""
    distractors = [p for p in pool if p.lower() != correct.lower()]
    random.shuffle(distractors)
    distractors = distractors[:3]
    while len(distractors) < 3:
        distractors.append("___")

    labels = ["a", "b", "c", "d"]
    correct_pos = random.randint(0, 3)
    options: list[QuestionOption] = []
    di = 0
    for i, label in enumerate(labels):
        if i == correct_pos:
            options.append(QuestionOption(label=label, text=correct))
        else:
            options.append(QuestionOption(label=label, text=distractors[di]))
            di += 1
    return options, labels[correct_pos]


def _resolve_question(
    q: FoundationQuestion, fib_pool: list[str]
) -> tuple[str, list[QuestionOption], str]:
    """Returns (effective_type, effective_options, correct_display)."""
    answer_key = (q.answer_key or "").strip()
    q_type = q.question_type
    opts = list(q.options or q.shared_options)
    correct_display = answer_key

    # ALL FIB → SCQ regardless of whether answer is numeric or word
    if q_type == "FIB":
        q_type = "SCQ"
        if answer_key:
            opts, correct_label = _build_fib_options(answer_key, fib_pool)
            correct_display = f"({correct_label})"
        else:
            correct_display = ""

    return q_type, opts, correct_display


def _clean_latex(text: str) -> str:
    text = re.sub(r"\$\$([^$]+)\$\$", r"\1", text)
    text = re.sub(r"\$([^$]+)\$", r"\1", text)
    text = re.sub(r"\\mathrm\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\text\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\rightarrow|\\longrightarrow|\\xrightarrow\{[^}]*\}", " -> ", text)
    text = re.sub(r"\\\s+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


SUBJECTIVE_TYPES = frozenset({
    "VSA", "SA", "LA", "DESCRIPTIVE", "TEXTBOOK", "PASSAGE", "CASE_STUDY",
})

OBJECTIVE_TYPES = frozenset({
    "MCQ", "MCQ_MULTI", "ASSERTION_REASON", "FIB", "SCQ", "TRUE_FALSE",
    "INTEGER", "MATCH",
})


def export_docx(doc: FoundationDocument, out_path: str | Path) -> Path:
    """Export objective-only question bank to .docx."""
    out_path = Path(out_path)
    docx = Document()

    # Title
    title = docx.add_heading(doc.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = docx.add_paragraph(f"Foundation {doc.subject} | Class {doc.class_level}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    objective_qs = [q for q in doc.questions if q.question_type not in SUBJECTIVE_TYPES]
    subjective_count = len(doc.questions) - len(objective_qs)

    if subjective_count:
        note = docx.add_paragraph(
            f"Note: {subjective_count} subjective question(s) excluded — see separate download."
        )
        note.runs[0].italic = True
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fib_pool = _build_fib_pool(doc)
    current_section = ""
    seq = 0

    for q in objective_qs:
        section_label = q.section

        if section_label != current_section:
            current_section = section_label
            docx.add_heading(q.section, level=2)

            if q.directions:
                p = docx.add_paragraph(q.directions)
                p.runs[0].italic = True

            if q.passage:
                docx.add_paragraph(f"Passage: {_clean_latex(q.passage)}")

        was_fib = q.question_type == "FIB"
        q_type, opts, correct_display = _resolve_question(q, fib_pool)
        seq += 1

        # Question stem
        stem_text = _clean_latex(q.stem.replace("\n", " "))
        p = docx.add_paragraph()
        run = p.add_run(f"Q{seq}. [{q_type}]  {stem_text}")
        run.bold = True

        # Options
        for opt in opts:
            docx.add_paragraph(
                f"  ({opt.label}) {_clean_latex(opt.text)}",
                style="List Bullet"
            )

        # Column A / B (matching type)
        if q.column_a:
            docx.add_paragraph("  Column I")
            for opt in q.column_a:
                docx.add_paragraph(f"    ({opt.label}) {_clean_latex(opt.text)}")
        if q.column_b:
            docx.add_paragraph("  Column II")
            for opt in q.column_b:
                docx.add_paragraph(f"    ({opt.label}) {_clean_latex(opt.text)}")

        # Correct answer
        if correct_display:
            p = docx.add_paragraph()
            run = p.add_run(f"► Correct Answer: {correct_display}")
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            run.bold = True

        # Explanation
        if q.explanation:
            src = " [LLM]" if q.answer_source == "llm" else ""
            p = docx.add_paragraph()
            run = p.add_run(f"★ Explanation{src}: {_clean_latex(q.explanation)}")
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
            run.italic = True

    docx.save(out_path)
    return out_path


def export_subjective_docx(doc: FoundationDocument, out_path: str | Path) -> Path:
    """Export subjective (non-convertible) questions to a separate .docx."""
    out_path = Path(out_path)
    subjective_qs = [q for q in doc.questions if q.question_type in SUBJECTIVE_TYPES]

    docx = Document()

    title = docx.add_heading(f"{doc.title} — Subjective Questions", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = docx.add_paragraph(f"Foundation {doc.subject} | Class {doc.class_level}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    note = docx.add_paragraph(
        f"Total: {len(subjective_qs)} subjective question(s). "
        "These could not be converted to MCQ/SCQ/INTEGER format."
    )
    note.runs[0].italic = True
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    current_section = ""
    seq = 0

    for q in subjective_qs:
        section_label = q.section

        if section_label != current_section:
            current_section = section_label
            docx.add_heading(q.section, level=2)

            if q.directions:
                p = docx.add_paragraph(q.directions)
                p.runs[0].italic = True

            if q.passage:
                docx.add_paragraph(f"Passage: {_clean_latex(q.passage)}")

        seq += 1
        stem_text = _clean_latex(q.stem.replace("\n", " "))
        p = docx.add_paragraph()
        run = p.add_run(f"Q{seq}. [{q.question_type}]  {stem_text}")
        run.bold = True

        if q.answer_key:
            p = docx.add_paragraph()
            run = p.add_run(f"► Answer: {_clean_latex(q.answer_key)}")
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            run.bold = True

        if q.explanation:
            src = " [LLM]" if q.answer_source == "llm" else ""
            p = docx.add_paragraph()
            run = p.add_run(f"★ Explanation{src}: {_clean_latex(q.explanation)}")
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
            run.italic = True

    docx.save(out_path)
    return out_path


def document_to_docx_bytes(doc: FoundationDocument) -> bytes:
    """Return the objective question bank .docx as raw bytes (for HTTP response)."""
    import io
    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)

    export_docx(doc, tmp_path)
    data = tmp_path.read_bytes()
    tmp_path.unlink(missing_ok=True)
    return data


def subjective_docx_bytes(doc: FoundationDocument) -> bytes:
    """Return the subjective questions .docx as raw bytes (for HTTP response)."""
    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)

    export_subjective_docx(doc, tmp_path)
    data = tmp_path.read_bytes()
    tmp_path.unlink(missing_ok=True)
    return data    while len(distractors) < 3:
        distractors.append("___")

    labels = ["a", "b", "c", "d"]
    correct_pos = random.randint(0, 3)
    options: list[QuestionOption] = []
    di = 0
    for i, label in enumerate(labels):
        if i == correct_pos:
            options.append(QuestionOption(label=label, text=correct))
        else:
            options.append(QuestionOption(label=label, text=distractors[di]))
            di += 1
    return options, labels[correct_pos]


def _resolve_question(
    q: FoundationQuestion, fib_pool: list[str]
) -> tuple[str, list[QuestionOption], str]:
    """Returns (effective_type, effective_options, correct_display)."""
    answer_key = (q.answer_key or "").strip()
    q_type = q.question_type
    opts = list(q.options or q.shared_options)
    correct_display = answer_key

    if q_type == "FIB" and answer_key and not _is_nonneg_numeric(answer_key):
        q_type = "SCQ"
        opts, correct_label = _build_fib_options(answer_key, fib_pool)
        correct_display = f"({correct_label})"

    return q_type, opts, correct_display


def _clean_latex(text: str) -> str:
    text = re.sub(r"\$\$([^$]+)\$\$", r"\1", text)
    text = re.sub(r"\$([^$]+)\$", r"\1", text)
    text = re.sub(r"\\mathrm\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\text\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\rightarrow|\\longrightarrow|\\xrightarrow\{[^}]*\}", " -> ", text)
    text = re.sub(r"\\\s+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def export_docx(doc: FoundationDocument, out_path: str | Path) -> Path:
    out_path = Path(out_path)
    docx = Document()

    # Title
    title = docx.add_heading(doc.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = docx.add_paragraph(f"Foundation {doc.subject} | Class {doc.class_level}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fib_pool = _build_fib_pool(doc)
    current_section = ""
    seq = 0

    for q in doc.questions:
        section_label = q.section

        if section_label != current_section:
            current_section = section_label
            docx.add_heading(q.section, level=2)

            if q.directions:
                p = docx.add_paragraph(q.directions)
                p.runs[0].italic = True

            if q.passage:
                docx.add_paragraph(f"Passage: {_clean_latex(q.passage)}")

        was_fib = q.question_type == "FIB"
        q_type, opts, correct_display = _resolve_question(q, fib_pool)
        seq += 1

        # Question stem
        stem_text = _clean_latex(q.stem.replace("\n", " "))
        p = docx.add_paragraph()
        run = p.add_run(f"Q{seq}. [{q_type}]  {stem_text}")
        run.bold = True

        # Options
        for opt in opts:
            docx.add_paragraph(
                f"  ({opt.label}) {_clean_latex(opt.text)}",
                style="List Bullet"
            )

        # Column A / B (matching type)
        if q.column_a:
            docx.add_paragraph("  Column I")
            for opt in q.column_a:
                docx.add_paragraph(f"    ({opt.label}) {_clean_latex(opt.text)}")
        if q.column_b:
            docx.add_paragraph("  Column II")
            for opt in q.column_b:
                docx.add_paragraph(f"    ({opt.label}) {_clean_latex(opt.text)}")

        # Correct answer
        if correct_display:
            p = docx.add_paragraph()
            run = p.add_run(f"► Correct Answer: {correct_display}")
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            run.bold = True

        # Explanation
        if q.explanation:
            src = " [LLM]" if q.answer_source == "llm" else ""
            p = docx.add_paragraph()
            run = p.add_run(f"★ Explanation{src}: {_clean_latex(q.explanation)}")
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
            run.italic = True

    docx.save(out_path)
    return out_path
